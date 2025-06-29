import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import docx

from src.extraction.base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """Extractor for DOCX files (medical reports, clinical summaries, etc.)."""
    
    def __init__(self):
        super().__init__()
        self.doc = None
        self.paragraphs = []
        self.tables = []
        self.date_pattern = re.compile(r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b')
        
        # For section detection
        self.medical_sections = [
            "assessment", "diagnosis", "medications", "allergies", "history",
            "physical examination", "lab results", "plan", "treatment", 
            "referral", "follow-up", "chief complaint", "review of systems"
        ]
    
    def _extract_metadata(self) -> Dict:
        """Extract metadata from the DOCX file."""
        metadata = {}
        
        # Extract date from filename
        file_date = self._extract_date_from_filename()
        if file_date:
            metadata["file_date"] = file_date
        
        # Get file creation and modification times
        stat = self.source_file.stat()
        metadata["creation_time"] = datetime.fromtimestamp(stat.st_ctime).isoformat()
        metadata["modification_time"] = datetime.fromtimestamp(stat.st_mtime).isoformat()
        
        # Basic file properties
        metadata["file_size"] = stat.st_size
        metadata["file_type"] = "docx"
        metadata["file_name"] = self.source_file.name
        
        # Extract document properties from DOCX
        try:
            self.doc = docx.Document(self.source_file)
            
            # Extract core properties if available
            core_props = self.doc.core_properties
            if core_props:
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["doc_created"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["doc_modified"] = core_props.modified.isoformat()
                if core_props.last_modified_by:
                    metadata["last_modified_by"] = core_props.last_modified_by
                    
            # Count paragraphs and tables
            metadata["paragraph_count"] = len(self.doc.paragraphs)
            metadata["table_count"] = len(self.doc.tables)
            
        except Exception as e:
            metadata["docx_metadata_error"] = str(e)
        
        return metadata
    
    def _extract_content(self) -> str:
        """Extract content from the DOCX file using python-docx."""
        try:
            if not self.doc:
                self.doc = docx.Document(self.source_file)
            
            # Extract text from paragraphs
            self.paragraphs = []
            paragraph_texts = []
            
            for para in self.doc.paragraphs:
                text = para.text.strip()
                if text:
                    self.paragraphs.append(text)
                    paragraph_texts.append(text)
            
            # Extract tables
            self.tables = []
            for table in self.doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_data):  # Only add non-empty rows
                        table_data.append(row_data)
                
                if table_data:  # Only add non-empty tables
                    self.tables.append(table_data)
                    
                    # Add table content to extracted text
                    paragraph_texts.append("\nTABLE:")
                    for row in table_data:
                        paragraph_texts.append(" | ".join(row))
            
            # Combine all text
            full_text = "\n\n".join(paragraph_texts)
            
            # Set confidence score
            if full_text and len(full_text) > 100:
                self.confidence_score = 1.0
            elif full_text:
                self.confidence_score = 0.8
            else:
                self.confidence_score = 0.5
            
            return full_text
            
        except Exception as e:
            self.confidence_score = 0.0
            return f"Error extracting content: {str(e)}"
    
    def extract_dates(self) -> Set[str]:
        """Extract all dates mentioned in the content."""
        if not self.content:
            return set()
            
        date_matches = self.date_pattern.findall(self.content)
        normalized_dates = set()
        
        for date_str in date_matches:
            try:
                # Try to parse and normalize the date
                if '/' in date_str:
                    parts = date_str.split('/')
                elif '-' in date_str:
                    parts = date_str.split('-')
                else:
                    continue
                    
                # Handle different date formats
                if len(parts) == 3:
                    if len(parts[2]) == 4:  # MM/DD/YYYY
                        month, day, year = parts
                    else:  # YYYY/MM/DD
                        year, month, day = parts
                        
                    # Make sure year is 4 digits
                    if len(year) == 2:
                        if int(year) > 50:  # Assume 19xx for years > 50
                            year = f"19{year}"
                        else:  # Assume 20xx for years <= 50
                            year = f"20{year}"
                            
                    normalized_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                    normalized_dates.add(normalized_date)
            except:
                continue
                
        return normalized_dates
    
    def extract_sections(self) -> Dict[str, str]:
        """Extract medical sections from paragraphs based on headings and formatting."""
        sections = {}
        
        if not self.doc or not self.paragraphs:
            return sections
        
        # Look for sections based on headings (often bold or larger font)
        current_section = None
        section_content = []
        
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this paragraph is a heading
            is_heading = False
            
            # Check if paragraph style indicates a heading
            if paragraph.style.name.startswith('Heading'):
                is_heading = True
            
            # Check if paragraph is bold (might be a section header)
            elif any(run.bold for run in paragraph.runs):
                is_heading = True
            
            # Check if text is a medical section indicator
            elif text.lower() in self.medical_sections or any(
                section in text.lower() for section in self.medical_sections):
                is_heading = True
            
            # If heading found, start a new section
            if is_heading:
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content)
                
                # Start new section
                current_section = text
                section_content = []
            elif current_section:
                section_content.append(text)
        
        # Save the last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        return sections
    
    def extract_tables(self) -> List[Dict]:
        """Extract and convert tables to structured format."""
        structured_tables = []
        
        if not self.doc or not self.tables:
            return structured_tables
        
        for i, table_data in enumerate(self.tables):
            if not table_data or not table_data[0]:
                continue
                
            # Create a structured representation with headers
            table_dict = {
                "id": i,
                "headers": table_data[0],
                "rows": []
            }
            
            # Add rows as dicts with header keys
            for row in table_data[1:]:
                if len(row) == len(table_dict["headers"]):
                    row_dict = {table_dict["headers"][j]: row[j] for j in range(len(row))}
                    table_dict["rows"].append(row_dict)
                else:
                    # Handle mismatched columns by using indices
                    table_dict["rows"].append(row)
            
            structured_tables.append(table_dict)
        
        return structured_tables
    
    def extract_headings(self) -> List[str]:
        """Extract all headings from the document."""
        headings = []
        
        if not self.doc:
            return headings
        
        for paragraph in self.doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                text = paragraph.text.strip()
                if text:
                    headings.append(text)
        
        return headings
    
    def extract_medical_terms(self) -> List[Tuple[str, int]]:
        """Find medical terms in the document with their paragraph position."""
        medical_terms = []
        
        if not self.content or not self.paragraphs:
            return medical_terms
        
        # Common medical term patterns
        term_patterns = [
            re.compile(r'\b[Dd]iagnosis\b'),
            re.compile(r'\b[Pp]rescribed\b'),
            re.compile(r'\b[Mm]edication\b'),
            re.compile(r'\b[Tt]reatment\b'),
            re.compile(r'\b[Ss]ymptoms?\b'),
            re.compile(r'\b[Tt]herapy\b'),
            re.compile(r'\b[Cc]ondition\b'),
            re.compile(r'\b[Pp]atient\b'),
            re.compile(r'\b[Dd]osage\b'),
            re.compile(r'\b[Tt]est\b'),
            re.compile(r'\b[Ll]ab\b'),
            re.compile(r'\b[Rr]esults?\b')
        ]
        
        # Search for each term in each paragraph
        for i, para in enumerate(self.paragraphs):
            for pattern in term_patterns:
                matches = pattern.finditer(para)
                for match in matches:
                    term = match.group(0)
                    # Store term and its paragraph index
                    medical_terms.append((term, i))
        
        return medical_terms 